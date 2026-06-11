"""Genere le rapport docx expliquant le calcul des 12 coefficients
et leur integration dans le notebook."""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path(__file__).parent / "Rapport_12_Coefficients.docx"


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # gris clair fond
    rpr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    rpr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_table_3cols(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def main():
    doc = Document()

    # Styles globaux
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ========================================================================
    # PAGE DE GARDE
    # ========================================================================
    title = doc.add_heading("Calcul des 12 coefficients de la formule v_features", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Methodes explorees, resultats compares, integration dans le notebook")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # ========================================================================
    # 1. RAPPEL DU CONTEXTE
    # ========================================================================
    add_heading(doc, "1. Contexte et formule", 1)

    add_paragraph(doc,
        "Le pipeline d'evaluation predit le score d'occlusion d'un visage par "
        "une formule lineaire ponderee. Pour chaque image, on extrait 6 features "
        "via segmentation (BiSeNet + SegFormer + 3DDFA-V2), puis on applique :"
    )

    add_code_block(doc,
        "pred = w_hair_bi      * hair_ratio_b\n"
        "     + w_hat_bi       * hat_ratio_b\n"
        "     + w_other_bi     * other_ratio_b\n"
        "     + w_hair_sf      * hair_ratio_s\n"
        "     + w_hat_sf       * hat_ratio_s\n"
        "     + w_other_bg_sf  * other_ratio_s\n"
    )

    add_paragraph(doc,
        "Avec 6 coefficients par genre (Femmes, Hommes), on a 12 coefficients au total a calibrer."
    )

    add_paragraph(doc,
        "Question : comment trouver ces 12 valeurs ?"
    )

    # ========================================================================
    # 2. APPROCHE PEDAGOGIQUE POUR LE RAPPORT
    # ========================================================================
    add_heading(doc, "2. Approche pedagogique pour le rapport", 1)

    add_paragraph(doc,
        "Nous avons construit un pipeline d'evaluation qui :",
    )
    for txt in [
        "Tire une image au hasard du dataset val",
        "Calcule les features par segmentation (BiSeNet + SegFormer + 3DDFA-V2)",
        "Predit un score d'occlusion par formule lineaire ponderee",
    ]:
        p = doc.add_paragraph(txt, style="List Number")

    add_paragraph(doc,
        "Lors de tirages successifs, nous avons observe des biais systematiques :",
    )
    for txt in [
        "Sur-prediction sur les images avec beaucoup de cheveux ou chapeaux",
        "Sous-prediction sur les images avec accessoires (lunettes, micros, mains)",
        "Distributions de prediction differentes selon le genre",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    add_paragraph(doc,
        "Pour corriger ces biais, nous avons mene une petite etude statistique sur "
        "~1% du dataset (150 images sur 15K val), en restant dans l'esprit "
        "training-free : aucun modele entraine, juste de la statistique simple."
    )

    add_heading(doc, "2.1 Methode : regression lineaire 1D par feature", 2)

    add_paragraph(doc,
        "Pour chaque feature et chaque genre, on calcule la pente de la regression "
        "lineaire 1D qui minimise la somme des carres des erreurs entre les points "
        "(feature, target) :"
    )

    add_code_block(doc,
        "slope = Sum(X * y) / Sum(X * X)\n"
    )

    add_paragraph(doc, "Demonstration en 3 lignes :", bold=True)
    add_code_block(doc,
        "Loss(slope) = Sum_i (slope * X_i - y_i)^2\n"
        "Derivee par rapport a slope :  2 * Sum_i X_i * (slope * X_i - y_i) = 0\n"
        "==> slope = Sum_i (X_i * y_i) / Sum_i X_i^2\n"
    )

    add_paragraph(doc,
        "C'est la formule de la pente de la droite y = slope * X qui passe au mieux "
        "par les points observes. Elle s'apprend en premiere annee de stats. "
        "Aucune matrice, pas de bibliotheque magique, juste des sommes."
    )

    add_heading(doc, "2.2 Resultats sur 1% du dataset (150 images, seed=42)", 2)

    headers = ["Feature", "Coeff F (gender=0)", "Coeff M (gender=1)"]
    rows = [
        ("hair_bi",     "+1.0015", "+1.1502"),
        ("hat_bi",      "+1.1588", "+0.9447"),
        ("other_bi",    "+0.5636", "+0.3350"),
        ("hair_sf",     "+0.8812", "+0.7504"),
        ("hat_sf",      "+1.1455", "+0.7633"),
        ("other_bg_sf", "+0.8151", "+0.3993"),
    ]
    add_table_3cols(doc, headers, rows)

    add_paragraph(doc,
        "Stabilite : recalcul avec 5 seeds differentes pour estimer la variance. "
        "Certains coefficients (hat_bi F, hat_sf F) sont peu stables car les "
        "chapeaux sont rares dans 45 images. Pour stabiliser, il faudrait un "
        "echantillon plus grand."
    )

    add_paragraph(doc,
        "Cette methode donne un score brief de ~0.062 sur les 15K val (vs 0.006 "
        "pour la methode optimale). Le facteur 10 d'ecart provient du fait que "
        "la regression 1D ignore les correlations entre features : hair_bi et "
        "hair_sf voient souvent les memes cheveux, donc chacun s'attribue tout "
        "le credit (double comptage)."
    )

    # ========================================================================
    # 3. METHODE UTILISEE POUR LA SUBMISSION
    # ========================================================================
    add_heading(doc, "3. Methode optimale (utilisee pour la submission)", 1)

    add_paragraph(doc,
        "Pour la submission finale, nous avons utilise une optimisation directe "
        "du score officiel IDEMIA par l'algorithme Nelder-Mead :"
    )

    add_code_block(doc,
        "from scipy.optimize import minimize\n"
        "\n"
        "def objective_brief(params, df):\n"
        "    F_w = {'hair_bi': params[0], 'hat_bi': params[1], ...}\n"
        "    M_w = {'hair_bi': params[6], 'hat_bi': params[7], ...}\n"
        "    df['pred'] = predict_v_features(df, F_w, M_w)\n"
        "    return reweighted_score(df, TEST_DISTRIBUTIONS['brief'])['score']\n"
        "\n"
        "res = minimize(objective_brief, x0, args=(df_val,),\n"
        "               method='Nelder-Mead',\n"
        "               options={'maxiter': 5000, 'xatol': 1e-5})\n"
    )

    add_paragraph(doc,
        "Nelder-Mead est un algorithme de simplex sans gradient : il evalue le score "
        "en 12 points (le simplex), bouge ces points pour descendre vers le minimum. "
        "Il converge en ~5000 iterations sur les 15K val et donne le meilleur score "
        "brief possible avec cette formule."
    )

    add_heading(doc, "3.1 Coefficients optimaux trouves", 2)

    headers = ["Feature", "Coeff F (gender=0)", "Coeff M (gender=1)"]
    rows = [
        ("hair_bi",     "+0.376", "+0.489"),
        ("hat_bi",      "+0.425", "+0.294"),
        ("other_bi",    "+0.478", "+0.210"),
        ("hair_sf",     "+0.619", "+0.484"),
        ("hat_sf",      "+0.902", "+0.609"),
        ("other_bg_sf", "+1.087", "+0.382"),
    ]
    add_table_3cols(doc, headers, rows)

    add_paragraph(doc,
        "Score brief obtenu sur 15K val : 0.00594 (= score officiel IDEMIA reweighted)."
    )

    # ========================================================================
    # 4. COMPARAISON COMPLETE
    # ========================================================================
    add_heading(doc, "4. Comparaison des methodes", 1)

    headers = ["Methode", "Brief score (15K val)", "vs Nelder-Mead"]
    rows = [
        ("Nelder-Mead (15K val)",                    "0.00594", "reference"),
        ("Regression lineaire ponderee WLS (sklearn)","0.00786", "+32%"),
        ("Regression 1D par feature (15K val)",      "0.07422", "+1150%"),
        ("Regression 1D par feature (1% = 150)",     "0.06238", "+950%"),
    ]
    add_table_3cols(doc, headers, rows)

    add_paragraph(doc,
        "Interpretation :",
        bold=True,
    )

    for txt in [
        "Nelder-Mead est le meilleur car il optimise DIRECTEMENT la metrique cible (brief score officiel).",
        "WLS sklearn est proche (+32%) car il prend en compte les correlations entre features, mais il optimise MSE et pas le brief.",
        "La regression 1D par feature est ~10x pire car elle ignore les correlations (double comptage hair_bi + hair_sf).",
        "Sur 1% du dataset, on obtient des coefficients legerement differents mais qualitativement similaires.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    # ========================================================================
    # 5. INTEGRATION DANS LE NOTEBOOK
    # ========================================================================
    add_heading(doc, "5. Integration des coefficients dans le notebook", 1)

    add_paragraph(doc,
        "Le notebook DataChallengeJulien_TF_with_v2_fallback.ipynb utilise les "
        "coefficients dans la cellule des constantes (cellule 35)."
    )

    add_heading(doc, "5.1 Localisation dans le notebook", 2)

    add_code_block(doc,
        '# Dans la cellule 35 :\n'
        '\n'
        'F_WEIGHTS = {"hair_bi": 0.376, "hat_bi": 0.425, "other_bi": 0.478,\n'
        '             "hair_sf": 0.619, "hat_sf": 0.902, "other_bg_sf": 1.087}\n'
        'M_WEIGHTS = {"hair_bi": 0.489, "hat_bi": 0.294, "other_bi": 0.210,\n'
        '             "hair_sf": 0.484, "hat_sf": 0.609, "other_bg_sf": 0.382}\n'
    )

    add_paragraph(doc,
        "Ces deux dictionnaires sont passes a la fonction apply_v2_fallback_with_dampening() "
        "dans l'orchestrateur occlusion_computation()."
    )

    add_heading(doc, "5.2 Pour changer de coefficients", 2)

    add_paragraph(doc,
        "Si Julien veut tester d'autres coefficients (par exemple ceux issus de la "
        "regression 1D pedagogique sur 1%), il suffit de remplacer les valeurs dans "
        "la cellule 35 :"
    )

    add_code_block(doc,
        '# Version pedagogique (regression 1D sur 1% du dataset) :\n'
        '\n'
        'F_WEIGHTS = {"hair_bi": 1.0015, "hat_bi": 1.1588, "other_bi": 0.5636,\n'
        '             "hair_sf": 0.8812, "hat_sf": 1.1455, "other_bg_sf": 0.8151}\n'
        'M_WEIGHTS = {"hair_bi": 1.1502, "hat_bi": 0.9447, "other_bi": 0.3350,\n'
        '             "hair_sf": 0.7504, "hat_sf": 0.7633, "other_bg_sf": 0.3993}\n'
    )

    add_paragraph(doc,
        "Le reste du notebook (sous-fonctions, orchestrateur, fallback) reste inchange. "
        "Seule la cellule 35 est modifiee."
    )

    # ========================================================================
    # 6. SCRIPTS DISPONIBLES
    # ========================================================================
    add_heading(doc, "6. Scripts disponibles dans pipeline_julien_integration/", 1)

    add_paragraph(doc, "Quatre scripts standalone pour recalculer les coefficients :")

    headers = ["Script", "Methode", "Brief obtenu"]
    rows = [
        ("optimize_12_coefficients.py",            "Nelder-Mead (15K val)",          "0.00594"),
        ("optimize_12_coefficients_simple.py",     "Reg lin ponderee (sklearn)",     "0.00786"),
        ("optimize_12_coefficients_pedagogique.py","Reg 1D par feature (15K val)",   "0.07422"),
        ("optimize_12_coefficients_1pct_dataset.py","Reg 1D sur 1% du dataset",      "0.06238"),
    ]
    add_table_3cols(doc, headers, rows)

    add_paragraph(doc,
        "Chacun produit un fichier JSON avec les 12 coefficients. Julien peut les "
        "relancer pour comparer ou explorer d'autres approches."
    )

    add_heading(doc, "6.1 Usage type", 2)

    add_code_block(doc,
        "# Methode optimale (utilisee pour la submission)\n"
        "python optimize_12_coefficients.py \\\n"
        "    --val-cache eval/cache/val_features.csv \\\n"
        "    --out coefficients.json\n"
        "\n"
        "# Methode pedagogique sur 1% du dataset (pour le rapport)\n"
        "python optimize_12_coefficients_1pct_dataset.py \\\n"
        "    --val-cache eval/cache/val_features.csv \\\n"
        "    --out coefficients_1pct.json\n"
    )

    # ========================================================================
    # 7. POUR LE RAPPORT
    # ========================================================================
    add_heading(doc, "7. Structure suggeree pour le rapport", 1)

    add_paragraph(doc,
        "Pour presenter la calibration dans le rapport, nous suggerons une "
        "progression pedagogique en 3 etapes :"
    )

    add_paragraph(doc, "Etape 1 - Approche statistique simple", bold=True)
    add_paragraph(doc,
        "On tire 1% du dataset (150 images). Pour chaque feature et chaque genre, "
        "on calcule la pente d'une regression lineaire 1D : slope = Sum(X*y) / Sum(X*X). "
        "Cela donne 12 coefficients en une formule mathematique pure, sans optimisation. "
        "Score brief : 0.062. On observe que la methode ignore les correlations -> "
        "limite."
    )

    add_paragraph(doc, "Etape 2 - Regression multi-feature ponderee", bold=True)
    add_paragraph(doc,
        "Pour corriger le double comptage, on resout le systeme entier : "
        "w = (X^T W X)^-1 X^T W y, avec W = matrice diagonale des poids 1/30 + target. "
        "Les correlations entre features sont prises en compte. Score brief : 0.008."
    )

    add_paragraph(doc, "Etape 3 - Optimisation directe du score officiel", bold=True)
    add_paragraph(doc,
        "Le score brief n'est pas une simple MSE : il pondere par bins et ajoute une "
        "penalite de fairness entre genres. Pour l'optimiser directement, on utilise "
        "Nelder-Mead (simplex sans gradient). Score brief : 0.006. C'est cette version "
        "qui est utilisee pour la submission."
    )

    add_paragraph(doc,
        "Cette progression raconte une histoire claire : on commence simple, on "
        "identifie les limites, on raffine progressivement. C'est le format classique "
        "d'un rapport de challenge."
    )

    # ========================================================================
    # FOOTER
    # ========================================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document genere automatiquement par generate_report_coefficients.py")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(str(OUT))
    print(f"Rapport genere : {OUT}")


if __name__ == "__main__":
    main()
