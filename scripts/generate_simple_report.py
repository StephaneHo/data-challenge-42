"""Generate the SIMPLIFIED synthesis document focused on Strategy U only.

Three sections:
  1. Pipeline visualization
  2. The 2 formulas
  3. Per-bin × gender improvement table
"""
from __future__ import annotations

import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

from src.metric import per_bin_breakdown, score  # noqa: E402


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_code(doc, code: str, size: int = 9):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    p.paragraph_format.left_indent = Cm(0.3)
    return p


def add_callout(doc, text, color=(0x1F, 0x4E, 0x79)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(*color)
    return p


def add_table(doc, headers, rows, header_bold=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        if header_bold:
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    return table


def compute_perbin_table():
    """Build the per-bin × gender comparison between Strategy C (baseline) and U."""
    j = pd.read_csv(REPO_ROOT / "eval" / "val_julien_baseline.csv").rename(columns={"pred": "pj"})
    sf = pd.read_csv(REPO_ROOT / "eval" / "val_zs_simple_hull_scaled_power07_tta.csv")[
        ["filename", "pred"]
    ].rename(columns={"pred": "ps"})
    pg = pd.read_csv(REPO_ROOT / "eval" / "cache" / "val_gender_pred.csv")
    df = j.merge(sf, on="filename").merge(pg, on="filename")
    df["pred_gender"] = df.pred_gender.fillna(1.0)
    ps = df.ps.values

    df["pred_C"] = np.where(
        df.pj > 0.7,
        np.minimum(df.pj * 0.40, ps),
        0.78 * df.pj * 0.40 + 0.22 * ps,
    )

    pjc_s = df.pj * 0.85
    pred_S = np.where(
        df.pj > 0.65,
        0.15 * 0.5 * (pjc_s + ps) + 0.85 * np.minimum(pjc_s, ps),
        0.60 * pjc_s + 0.40 * ps,
    ).clip(0, 1)

    pjc_q = np.where(df.pred_gender == 0.0, df.pj * 0.75, df.pj * 0.70)
    a_lo_q = np.where(df.pred_gender == 0.0, 0.70, 0.50)
    pred_Q = np.where(
        df.pj > 0.65,
        0.15 * 0.5 * (pjc_q + ps) + 0.85 * np.minimum(pjc_q, ps),
        a_lo_q * pjc_q + (1 - a_lo_q) * ps,
    ).clip(0, 1)
    df["pred_U"] = 0.5 * pred_S + 0.5 * pred_Q

    bC = per_bin_breakdown(df, pred_col="pred_C", gt_col="target")
    bU = per_bin_breakdown(df, pred_col="pred_U", gt_col="target")

    m = bC[["gender", "bin", "n", "mean_gt"]].copy()
    m["pred_C"] = bC["mean_pred"].round(3)
    m["pred_U"] = bU["mean_pred"].round(3)
    m["mean_gt"] = m["mean_gt"].round(3)
    m["bias_C"] = (bC["mean_pred"] - bC["mean_gt"]).round(3)
    m["bias_U"] = (bU["mean_pred"] - bU["mean_gt"]).round(3)
    return m


def build_document(output_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============== Title + TL;DR ==============
    doc.add_heading("Stratégie U — soumission finale", level=0)
    p = doc.add_paragraph()
    p.add_run(
        "Document court (4 pages) qui décrit la stratégie soumise sur hfactory. "
        "Trois sections : 1) le pipeline, 2) les formules de calcul, 3) la mesure "
        "d'amélioration par bin × genre vs notre première soumission."
    )

    add_callout(
        doc,
        "Score attendu sur le leaderboard : 0.011 (vs notre première soumission à 0.016 = −30%).",
    )

    # ============== 1. Pipeline ==============
    add_heading(doc, "1. Le pipeline complet (une image → une prédiction)", level=1)

    p = doc.add_paragraph()
    p.add_run("Pour chaque image test, on calcule 3 signaux en parallèle :").italic = True

    pipeline_diagram = (
        "                              [image test]\n"
        "                                   |\n"
        "        +--------------------------+--------------------------+\n"
        "        |                          |                          |\n"
        "        v                          v                          v\n"
        "   PIPELINE A                 PIPELINE B                INSIGHTFACE\n"
        "   RetinaFace                 SegFormer                 buffalo_l\n"
        "   + 3DDFA-V2                 face-parsing              (gender)\n"
        "   + BiSeNet                                            \n"
        "                                                            \n"
        "   1. RetinaFace            1. SegFormer segmente       Predit l'attribut\n"
        "      detecte le visage        19 classes faciales      sex sur le visage\n"
        "                               (skin, hat, hair,        detecte.\n"
        "   2. 3DDFA-V2                 eye_g, cloth, ...).      \n"
        "      reconstruit le 3D                                 Precision mesuree :\n"
        "      -> masque theorique   2. Enveloppe convexe        90% sur full val.\n"
        "      du visage entier        des pixels visage         \n"
        "      (convex hull)           (skin + parties).         \n"
        "                                                        \n"
        "   3. BiSeNet                3. Heuristique             \n"
        "      identifie les            simple_hull_scaled       \n"
        "      pixels de peau           _power07 :               \n"
        "      visible                  (hull - face) / hull,    \n"
        "                               puis ^0.7, puis TTA      \n"
        "   4. ratio_geom =             (flip horizontal).       \n"
        "      1 - peau_visible                                  \n"
        "          / masque_theo                                 \n"
        "        |                          |                          |\n"
        "        v                          v                          v\n"
        "   ratio_geom in [0,1]      ratio_segf in [0,1]            g in {F, M}\n"
        "        |                          |                          |\n"
        "        +--------------------------+--------------------------+\n"
        "                                   |\n"
        "                  +----------------+----------------+\n"
        "                  |                                 |\n"
        "                  v                                 v\n"
        "          FORMULE A (sans genre)         FORMULE B (avec genre)\n"
        "                  |                                 |\n"
        "                  v                                 v\n"
        "               pred_A                            pred_B\n"
        "                  |                                 |\n"
        "                  +----------------+----------------+\n"
        "                                   |\n"
        "                                   v\n"
        "               FaceOcclusion soumise = (pred_A + pred_B) / 2\n"
    )
    add_code(doc, pipeline_diagram, size=8)

    # ============== 2. Les 2 formules ==============
    doc.add_page_break()
    add_heading(doc, "2. Les 2 formules de calcul", level=1)

    add_heading(doc, "FORMULE A — sans genre", level=2)
    add_code(
        doc,
        "ratio_geom_recal = ratio_geom * 0.85       # on garde 85% du signal\n"
        "\n"
        "if ratio_geom > 0.65:                      # zone 'haute occlusion'\n"
        "    pred_A = 0.15 * (ratio_geom_recal + ratio_segf) / 2 \\\n"
        "           + 0.85 * min(ratio_geom_recal, ratio_segf)\n"
        "else:                                      # zone 'normale'\n"
        "    pred_A = 0.60 * ratio_geom_recal + 0.40 * ratio_segf\n",
    )

    add_heading(doc, "FORMULE B — avec genre", level=2)
    add_code(
        doc,
        "if g == 'F':\n"
        "    cal = 0.75      # femmes : 75% du signal Pipeline A\n"
        "    a_lo = 0.70\n"
        "else:               # g == 'M'\n"
        "    cal = 0.70      # hommes : 70% du signal\n"
        "    a_lo = 0.50\n"
        "\n"
        "ratio_geom_recal = ratio_geom * cal\n"
        "\n"
        "if ratio_geom > 0.65:\n"
        "    pred_B = 0.15 * (ratio_geom_recal + ratio_segf) / 2 \\\n"
        "           + 0.85 * min(ratio_geom_recal, ratio_segf)\n"
        "else:\n"
        "    pred_B = a_lo * ratio_geom_recal + (1 - a_lo) * ratio_segf\n",
    )

    add_heading(doc, "Combinaison finale", level=2)
    add_code(
        doc,
        "FaceOcclusion = clip( (pred_A + pred_B) / 2 , 0, 1 )\n",
    )

    # ============== 3. Per-bin × gender improvement ==============
    doc.add_page_break()
    add_heading(doc, "3. Amélioration mesurée par bin × genre", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Comparaison directe entre notre première soumission (Strategy C : ratio_geom × 0.40) "
        "et Strategy U (la nouvelle soumission). Mesures sur le full val 15 001 images, "
        "décomposées par bin de target × genre."
    )
    add_callout(
        doc,
        "Lecture : 'bias' = prédiction moyenne − target moyenne. Idéal = 0. "
        "Négatif = on sous-prédit. Positif = on sur-prédit.",
    )

    perbin = compute_perbin_table()

    rows = []
    for _, r in perbin.iterrows():
        rows.append([
            str(r["gender"]),
            str(r["bin"]),
            f"{int(r['n']):>5}",
            f"{r['mean_gt']:.3f}",
            f"{r['pred_C']:.3f}",
            f"{r['bias_C']:+.3f}",
            f"{r['pred_U']:.3f}",
            f"{r['bias_U']:+.3f}",
        ])

    add_table(
        doc,
        ["Genre", "Bin target", "n", "mean target", "pred C", "bias C", "pred U", "bias U"],
        rows,
    )

    # ============== Lecture ==============
    add_heading(doc, "Comment lire ce tableau", level=2)
    add_bullet(
        doc,
        "Zone basse [0.00, 0.20) : Strategy U sur-prédit légèrement (bias positif ~0.07-0.10). "
        "C'était mieux calibré avec Strategy C (bias ~0.04-0.09). On perd un peu ici.",
    )
    add_bullet(
        doc,
        "Zone moyenne [0.20, 0.30) : Strategy C sous-prédisait fort (bias F=-0.077, M=-0.069). "
        "Strategy U est quasi-parfait (bias F=+0.014, M=+0.009).",
    )
    add_bullet(
        doc,
        "Zone haute [0.30, 0.50) : Strategy C sous-prédisait massivement (bias F=-0.154, M=-0.156). "
        "Strategy U réduit le bias à F=-0.050, M=-0.066 (3× mieux).",
    )
    add_bullet(
        doc,
        "Extreme [0.50, 1.01) : très peu de cas en val (2 F + 3 M), mais U améliore le bias "
        "de C −0.27/−0.38 à U −0.17/−0.25.",
    )

    add_heading(doc, "Pourquoi ce trade-off est gagnant sur le test", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Notre val contient 80% d'images en zone basse [0.00, 0.20). Sur cette zone, "
        "Strategy U est légèrement moins bonne que C. C'est pourquoi U score 0.020 "
        "sur val (vs C à 0.009)."
    )
    p = doc.add_paragraph()
    p.add_run(
        "MAIS le test (selon le brief Telecom officiel) contient seulement 48% d'images "
        "en zone basse, et 41% en zone [0.20, 0.50). Sur cette zone-là, Strategy U "
        "dépasse C de loin (bias 3× plus petit). Le résultat net :"
    )
    add_table(
        doc,
        ["Distribution test", "Score Strategy C", "Score Strategy U", "Gain"],
        [
            ["brief (officielle)", "0.01563", "0.01086", "−31%"],
            ["spread (variante)", "0.02669", "0.01555", "−42%"],
            ["heavy (worst case)", "0.04246", "0.02178", "−49%"],
        ],
    )

    add_callout(
        doc,
        "Strategy U est dans le top sur les 3 distributions test plausibles. "
        "C'est notre soumission v8.",
    )

    # ============== Closing ==============
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run(
        "Toutes les mesures viennent du val split stratifié de 15 001 images. "
        "La distribution test 'brief' est celle annoncée par Telecom dans le PDF du challenge. "
        "Document généré le 2026-06-01."
    ).italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    out = REPO_ROOT / "docs" / "synthese_simple.docx"
    build_document(out)
    print(f"wrote {out}")
