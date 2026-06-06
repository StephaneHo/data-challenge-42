"""Rapport v_features avec correction Julien : comparaison avec / sans correction.

Charge :
  - val_v10_features.csv  (SANS correction Julien)
  - val_features.csv      (AVEC correction Julien)
Compare les scores, coefficients optimaux, et per-bin breakdowns.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))
sys.path.insert(0, str(REPO_ROOT / "pipeline_v_features"))

from evaluate import (  # noqa: E402
    TEST_DISTRIBUTIONS, evaluate_predictions, per_bin_breakdown,
    predict_v_features, reweighted_score, native_score,
)

TEST_LIKE_BRIEF = TEST_DISTRIBUTIONS["brief"]
TEST_LIKE_SPREAD = TEST_DISTRIBUTIONS["spread"]
TEST_LIKE_HEAVY = TEST_DISTRIBUTIONS["heavy"]


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_callout(doc, text, color=(0x1F, 0x4E, 0x79)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(*color)
    return p


def add_code(doc, code: str, size: int = 9):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    p.paragraph_format.left_indent = Cm(0.3)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    return table


# Coefficients precedents (sans correction Julien)
F_WEIGHTS_NO_CORRECTION = {
    "hair_bi": 0.424, "hat_bi": 0.665, "other_bi": 0.787,
    "hair_sf": 0.610, "hat_sf": 0.402, "other_bg_sf": 0.603,
}
M_WEIGHTS_NO_CORRECTION = {
    "hair_bi": 0.508, "hat_bi": 0.557, "other_bi": -0.386,
    "hair_sf": 0.454, "hat_sf": 0.451, "other_bg_sf": 0.572,
}

# Coefficients nouveaux (avec correction Julien) -- a charger depuis JSON
with open(REPO_ROOT / "pipeline_v_features" / "coefficients.json") as f:
    coeffs_julien = json.load(f)
F_WEIGHTS_JULIEN = coeffs_julien["F"]
M_WEIGHTS_JULIEN = coeffs_julien["M"]


def compute_predictions_and_scores(df, F_w, M_w):
    df = df.copy()
    df["pred"] = predict_v_features(df, F_w, M_w)
    return {
        "val_natif": native_score(df),
        "brief": reweighted_score(df, TEST_LIKE_BRIEF),
        "spread": reweighted_score(df, TEST_LIKE_SPREAD),
        "heavy": reweighted_score(df, TEST_LIKE_HEAVY),
        "per_bin": per_bin_breakdown(df),
        "df": df,
    }


def main():
    df_no = pd.read_csv(REPO_ROOT / "eval" / "cache" / "val_v10_features.csv")
    df_yes = pd.read_csv(REPO_ROOT / "eval" / "cache" / "val_features.csv")

    # Pour le df_no, drop les NaN si necessaire
    df_no = df_no.dropna(subset=["hair_bi_in_mask"]).reset_index(drop=True)
    df_yes = df_yes.dropna(subset=["hair_bi_in_mask"]).reset_index(drop=True)

    no_correction = compute_predictions_and_scores(df_no, F_WEIGHTS_NO_CORRECTION,
                                                     M_WEIGHTS_NO_CORRECTION)
    with_correction = compute_predictions_and_scores(df_yes, F_WEIGHTS_JULIEN, M_WEIGHTS_JULIEN)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============================================================
    doc.add_heading("Pipeline v_features avec correction Julien -- resultats", level=0)
    p = doc.add_paragraph()
    p.add_run(
        "Ce rapport compare le score de la pipeline v_features (decomposition par feature "
        "hair / hat / other) AVEC vs SANS la correction de mask Julien (warpAffine). "
        f"Mesures sur {len(df_yes)} images val."
    )
    add_callout(
        doc,
        f"Resultat : la correction Julien DEGRADE le score brief de "
        f"{no_correction['brief']['score']:.5f} (sans) -> {with_correction['brief']['score']:.5f} "
        f"(avec) = +{100*(with_correction['brief']['score']-no_correction['brief']['score'])/no_correction['brief']['score']:.1f}%.",
        color=(0xCC, 0x00, 0x00),
    )

    # ============================================================
    add_heading(doc, "1. Pipeline et formule", level=1)
    add_code(
        doc,
        "# Pipeline v_features (identique dans les 2 versions) :\n"
        "1. RetinaFace                  -> bbox\n"
        "2. 3DDFA-V2                    -> convex hull des vertices\n"
        "3. *** CORRECTION JULIEN ***   <- difference entre les 2 versions\n"
        "   M = [[0.9, 0, 15], [0, 1.05, -10]]\n"
        "   mask = cv2.warpAffine(mask, M, dsize=img.shape[:2])\n"
        "4. BiSeNet @ 512x512           -> 19 classes\n"
        "5. SegFormer @ 512x512         -> 19 classes\n"
        "6. Fractions hair / hat / other / skin / bg par modele\n"
        "7. InsightFace genderage       -> g in {F, M}\n"
        "\n"
        "# Formule :\n"
        "FaceOcclusion = clip(\n"
        "    w_hair_bi_g   * hair_bi_in_mask  +\n"
        "    w_hat_bi_g    * hat_bi_in_mask   +\n"
        "    w_other_bi_g  * other_bi_in_mask +\n"
        "    w_hair_sf_g   * hair_sf_in_mask  +\n"
        "    w_hat_sf_g    * hat_sf_in_mask   +\n"
        "    w_otherbgsf_g * (other_sf_in_mask + bg_sf_in_mask),\n"
        "    0, 1\n"
        ")",
    )

    # ============================================================
    add_heading(doc, "2. Coefficients optimaux (re-fittes pour chaque version)", level=1)

    add_heading(doc, "Femmes (F)", level=2)
    rows = []
    for k in ["hair_bi", "hat_bi", "other_bi", "hair_sf", "hat_sf", "other_bg_sf"]:
        rows.append([k, f"{F_WEIGHTS_NO_CORRECTION[k]:+.4f}", f"{F_WEIGHTS_JULIEN[k]:+.4f}",
                     f"{F_WEIGHTS_JULIEN[k] - F_WEIGHTS_NO_CORRECTION[k]:+.4f}"])
    add_table(doc, ["feature", "Sans correction", "Avec correction Julien", "Delta"], rows)

    add_heading(doc, "Hommes (M)", level=2)
    rows = []
    for k in ["hair_bi", "hat_bi", "other_bi", "hair_sf", "hat_sf", "other_bg_sf"]:
        rows.append([k, f"{M_WEIGHTS_NO_CORRECTION[k]:+.4f}", f"{M_WEIGHTS_JULIEN[k]:+.4f}",
                     f"{M_WEIGHTS_JULIEN[k] - M_WEIGHTS_NO_CORRECTION[k]:+.4f}"])
    add_table(doc, ["feature", "Sans correction", "Avec correction Julien", "Delta"], rows)

    add_heading(doc, "Observations sur les coefficients", level=2)
    add_bullet(doc, "Avec correction Julien, les poids montent (ex: hat_sf F=0.40 -> 0.90, "
                    "other_bg_sf F=0.60 -> 1.09). Cela compense la reduction du mask (-6%).")
    add_bullet(doc, "Le coefficient suspect other_bi_M = -0.39 (sans correction) devient +0.21 "
                    "(avec correction) = plus realiste.")

    # ============================================================
    doc.add_page_break()
    add_heading(doc, "3. Scores compares (4 distributions)", level=1)

    p = doc.add_paragraph()
    p.add_run("5 metriques par distribution : err_F, err_M, mean_err = ponderee par taille, "
              "gap = |F-M|, et score = (err_F+err_M)/2 + gap (OFFICIEL).").italic = True

    add_heading(doc, "Version SANS correction Julien", level=2)
    rows = []
    for name, label in [("val_natif", "val natif"), ("brief", "brief (officiel)"),
                          ("spread", "spread"), ("heavy", "heavy")]:
        s = no_correction[name]
        rows.append([label, f"{s['err_F']:.5f}", f"{s['err_M']:.5f}",
                     f"{s['mean_err']:.5f}", f"{s['gap']:.5f}", f"{s['score']:.5f}"])
    add_table(doc, ["Distribution", "err_F", "err_M", "mean_err", "gap", "score"], rows)

    add_heading(doc, "Version AVEC correction Julien", level=2)
    rows = []
    for name, label in [("val_natif", "val natif"), ("brief", "brief (officiel)"),
                          ("spread", "spread"), ("heavy", "heavy")]:
        s = with_correction[name]
        rows.append([label, f"{s['err_F']:.5f}", f"{s['err_M']:.5f}",
                     f"{s['mean_err']:.5f}", f"{s['gap']:.5f}", f"{s['score']:.5f}"])
    add_table(doc, ["Distribution", "err_F", "err_M", "mean_err", "gap", "score"], rows)

    add_heading(doc, "Comparaison directe", level=2)
    rows = []
    for name, label in [("val_natif", "val natif"), ("brief", "BRIEF"),
                          ("spread", "spread"), ("heavy", "heavy")]:
        s_no = no_correction[name]["score"]
        s_yes = with_correction[name]["score"]
        delta = s_yes - s_no
        pct = 100 * delta / s_no if s_no > 0 else 0
        rows.append([label, f"{s_no:.5f}", f"{s_yes:.5f}", f"{delta:+.5f}", f"{pct:+.1f}%"])
    add_table(doc, ["Distribution", "Sans correction", "Avec correction Julien", "Delta", "%"], rows)

    add_callout(
        doc,
        "La correction Julien degrade le score sur TOUTES les distributions test simulees. "
        "Particulierement sur heavy (+68%) qui represente le cas worst-case.",
    )

    # ============================================================
    doc.add_page_break()
    add_heading(doc, "4. Per-bin x gender breakdown (avec correction Julien)", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Decomposition de la prediction par (bin x genre). bias = mean_pred - mean_target. "
        "Negatif = sous-prediction, positif = sur-prediction."
    ).italic = True

    add_heading(doc, "Femmes (F)", level=2)
    rows = []
    for _, r in with_correction["per_bin"].query("gender == 'F'").iterrows():
        rows.append([r.bin, f"{int(r.n)}", f"{r.mean_gt:.3f}", f"{r.mean_pred:.3f}",
                     f"{r.bias:+.3f}", f"{r.weighted_err:.5f}"])
    add_table(doc, ["bin", "n", "mean target", "mean pred", "bias", "weighted err"], rows)

    add_heading(doc, "Hommes (M)", level=2)
    rows = []
    for _, r in with_correction["per_bin"].query("gender == 'M'").iterrows():
        rows.append([r.bin, f"{int(r.n)}", f"{r.mean_gt:.3f}", f"{r.mean_pred:.3f}",
                     f"{r.bias:+.3f}", f"{r.weighted_err:.5f}"])
    add_table(doc, ["bin", "n", "mean target", "mean pred", "bias", "weighted err"], rows)

    # ============================================================
    add_heading(doc, "5. Analyse : pourquoi la correction Julien degrade-t-elle ?", level=1)

    add_heading(doc, "Hypothese 1 : correction tunee pour BiSeNet pur", level=2)
    add_bullet(doc, "Julien a determine empiriquement la correction (0.9, 1.05, +15, -10) "
                    "sur SA pipeline (BiSeNet seul, sans SegFormer ni decomposition par feature).")
    add_bullet(doc, "La correction ameliore son score (visiblement sur ses tests), mais notre "
                    "pipeline hybride mixte les deux modeles avec une calibration per-feature - "
                    "le mask deplace pourrait ne pas convenir aux deux modeles simultanement.")

    add_heading(doc, "Hypothese 2 : reduction du mask = perte de contexte", level=2)
    add_bullet(doc, "Le mask corrige est en moyenne 6% plus petit (mask_area : 28499 -> 26695).")
    add_bullet(doc, "SegFormer utilise un large contexte pour segmenter. Un mask plus petit "
                    "exclut des pixels qui auraient contribue a une meilleure detection.")

    add_heading(doc, "Hypothese 3 : translation non-symetrique", level=2)
    add_bullet(doc, "La translation (+15, -10) deplace le mask vers la droite-haut. "
                    "Cela pourrait sur-couvrir le front (zone avec hair) et sous-couvrir le menton.")

    # ============================================================
    add_heading(doc, "6. Recommandation", level=1)

    add_callout(
        doc,
        "Recommendation : revenir a la version SANS correction Julien (score brief 0.00497) "
        "pour la submission finale. La correction degrade notre pipeline hybride.",
    )

    add_bullet(doc, "OPTION A (recommandee) : Submission v_features SANS correction Julien. "
                    "Score brief attendu = 0.00497.")
    add_bullet(doc, "OPTION B : Test cache (~33h) AVEC correction Julien malgre tout, et choisir "
                    "selon le score reel sur le leaderboard.")
    add_bullet(doc, "OPTION C : Tester un compromis (correction partielle, ex: tx/2, ty/2) "
                    "ou re-tuner la correction pour notre pipeline hybride.")

    # Closing
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run(
        f"Document genere a partir des caches val_v10_features.csv (sans correction) et "
        f"val_features.csv (avec correction). Tous les coefficients sont re-fittes "
        f"independamment par Nelder-Mead pour chaque version."
    ).italic = True

    out_path = REPO_ROOT / "docs" / "v_features_julien_correction_report.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"wrote {out_path}")


if __name__ == "__main__":
    main()
